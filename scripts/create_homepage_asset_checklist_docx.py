from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUTPUT_PATH = r"C:\Users\sunny\Desktop\YIYI 网站重建\yiyimeshbeltconveyor-nextjs\Homepage_Asset_Checklist_YIYI_Mesh_Belt.docx"


ROWS = [
    ("1", "Hero 首屏背景", "Hero Video", "1", "机器人焊接 / 网带运行 / 网带与链轮啮合", "16:9 横版", "必需", "时长建议 8–20 秒"),
    ("2", "Hero 首屏背景", "Hero Poster Image", "1", "视频封面图，同场景最佳", "16:9 横版", "必需", "视频加载失败时显示"),
    ("3", "产品旗舰卡", "Spiral Freezer Belt", "1", "螺旋冷冻带主图或运行图", "4:3 或 16:10", "必需", "首页最大产品卡"),
    ("4", "产品重点卡", "Self-Stacker Belt", "1", "自堆叠网带产品图", "4:3", "必需", "第二层重点卡"),
    ("5", "产品重点卡", "Eye Link Belt", "1", "眼链网带产品图", "4:3", "必需", "第二层重点卡"),
    ("6", "产品补充卡", "Side Driven Belt", "1", "侧驱网带或相关系统图", "4:3", "建议", "现在可先用占位，后面替换"),
    ("7", "产品补充卡", "Flat Wire Honeycomb Belt", "1", "扁丝蜂窝带产品图", "4:3", "必需", "首页第三排产品卡"),
    ("8", "产品补充卡", "Drive Sprockets", "1", "链轮主图 / 配套图", "4:3", "必需", "突出 matched drive"),
    ("9", "制造信任模块", "Factory Overview", "1", "工厂车间全景 / 产线全景", "16:9 横版", "建议", "用于制造实力感"),
    ("10", "制造信任模块", "Robotic Welding", "1", "机器人焊接近景", "16:9 横版", "必需", "最适合首页信任段"),
    ("11", "制造信任模块", "In-House Wire Drawing", "1", "拉丝 / 原材处理 / 线材工艺", "16:9 横版", "建议", "配合正文卖点"),
    ("12", "制造信任模块", "Final Inspection / QC", "1", "质检、尺寸检查、鼓轮检查", "16:9 横版", "建议", "增强可靠性"),
    ("13", "出货与交付模块", "Export Packing", "1", "木箱、包装、唛头、装柜前", "16:9 横版", "建议", "配合 export-ready delivery"),
    ("14", "访厂/信任模块", "Factory Visit / Reception", "1", "接待、访厂、客户参观", "16:9 横版", "建议", "配合 Watch Factory Tour"),
    ("15", "案例证明模块", "Installation / Project Proof 01", "1", "客户现场安装图", "16:9 横版", "建议", "真实案例优先"),
    ("16", "案例证明模块", "Installation / Project Proof 02", "1", "设备配套 / 现场运行图", "16:9 横版", "建议", "可替换成出口/合作图"),
]


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


document = Document()

title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("YIYI Mesh Belt Homepage Asset Checklist")
run.bold = True
run.font.size = None

subtitle = document.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("首页素材清单表").italic = True

document.add_paragraph("")

table = document.add_table(rows=1, cols=8)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"

headers = ["序号", "模块位置", "素材名称", "数量", "建议内容", "建议比例", "是否必需", "备注"]
for cell, header in zip(table.rows[0].cells, headers):
    cell.text = header

for row in ROWS:
    cells = table.add_row().cells
    for idx, value in enumerate(row):
        cells[idx].text = value

document.add_paragraph("")
document.add_paragraph("最小可开工素材包")
for item in [
    "Hero Video",
    "Hero Poster Image",
    "Spiral Freezer Belt",
    "Self-Stacker Belt",
    "Eye Link Belt",
    "Flat Wire Honeycomb Belt",
    "Drive Sprockets",
    "Robotic Welding",
]:
    add_bullet(document, item)

document.add_paragraph("")
document.add_paragraph("建议文件名")
for item in [
    "home-hero-video.mp4",
    "home-hero-poster.jpg",
    "home-product-spiral-freezer.jpg",
    "home-product-self-stacker.jpg",
    "home-product-eye-link.jpg",
    "home-product-side-driven.jpg",
    "home-product-flat-wire.jpg",
    "home-product-sprocket.jpg",
    "home-factory-overview.jpg",
    "home-robotic-welding.jpg",
    "home-wire-drawing.jpg",
    "home-final-inspection.jpg",
    "home-export-packing.jpg",
    "home-factory-visit.jpg",
    "home-case-installation-01.jpg",
    "home-case-installation-02.jpg",
]:
    add_bullet(document, item)

document.add_paragraph("")
document.add_paragraph("素材标准")
for item in [
    "宽度尽量不低于 1600px",
    "横版优先",
    "真实工厂和真实产品优先",
    "避免过强滤镜和图库感",
    "能看到金属纹理、焊点、链轮啮合细节会更好",
]:
    add_bullet(document, item)

document.save(OUTPUT_PATH)
print(OUTPUT_PATH)
